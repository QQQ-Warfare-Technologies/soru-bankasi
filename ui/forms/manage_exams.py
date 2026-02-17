from PySide6.QtWidgets import QDialog, QMessageBox, QTableWidgetItem, QHeaderView, QApplication, QFileDialog
from PySide6.QtCore import Qt, QByteArray, QBuffer, QIODevice, QMimeData
from PySide6.QtGui import QPixmap
# Kendi UI dosyanın adını buraya yaz
from ui.forms.manage_exams_ui import Ui_Dialog
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


class ManageExamsDialog(QDialog):
    def __init__(self, db_manager, parent=None):
        super().__init__(parent)
        self.ui = Ui_Dialog()
        self.ui.setupUi(self)
        self.db = db_manager
        self.onizlenen_resim_yolu = None
        self.setWindowTitle("Sınav Kağıdı Hazırlama Modülü")
        self.resize(1000, 700)  # Geniş bir pencere olsun

        # --- 1. ARAYÜZ BAŞLANGIÇ AYARLARI ---
        # Yeni sınav ekleme alanlarını başlangıçta gizle
        self.ui.txt_yeni_sinav_adi.setVisible(False)
        self.ui.btn_sinav_kaydet.setVisible(False)
        self.ui.lbl_yeni_sinav.setVisible(False)
        # --- 2. BUTON BAĞLANTILARI ---
        # "+" butonuna basıldığında tetiklenecek fonksiyon
        self.ui.btn_sinav_arti.clicked.connect(
            self.yeni_sinav_alani_degistir)
        self.ui.btn_sinav_kaydet.clicked.connect(self.yeni_sinav_kaydet)
        self.ui.btn_sorulari_kaydet.clicked.connect(self.sinavi_tamamen_kaydet)
        self.ui.tree_filtre.itemChanged.connect(self.tabloyu_guncelle)
        self.ui.tree_filtre.itemClicked.connect(
            self.parent().tree_item_tiklandi)
        self.ui.table_all_questions.cellClicked.connect(
            self.havuz_satir_tiklandi)
        self.ui.table_exam_questions.cellClicked.connect(
            self.havuz_satir_tiklandi)
        self.ui.btn_ekle_sinava.clicked.connect(self.sinava_soru_ekle)
        self.ui.btn_cikar.clicked.connect(self.sinavdan_soru_cikar)
        self.ui.btn_soru_kopyala.clicked.connect(self.soru_kopyala)
        self.ui.btn_word_olustur.clicked.connect(self.word_dosyasi_olustur)
        self.ui.btn_sinav_sil.clicked.connect(self.sinav_sil)

        
        self.ui.table_exam_questions.setColumnCount(2)
        self.ui.table_exam_questions.setHorizontalHeaderLabels(
            ["ID", "Soru Metni"])
        self.ui.table_exam_questions.setColumnHidden(0, True)

        self.ui.cb_sinavlar.currentIndexChanged.connect(
            self.secili_sinav_sorularini_yukle)
        # --- 3. VERİ YÜKLEME ---
        self.sinavlari_comboboxa_yukle()
        self.parent().sol_menuyu_yukle(hedef_agac=self.ui.tree_filtre)

        # tablo ayarları
        self.ui.table_all_questions.setWordWrap(False)
        self.ui.table_exam_questions.setWordWrap(False)

        # 2. Yatay kaydırma çubuklarını "İhtiyaç Olduğunda Göster" olarak ayarla
        self.ui.table_all_questions.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.ui.table_exam_questions.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        # 3. Sütunların içeriğe göre sınırsız genişlemesine izin ver (Kritik ayar)
        self.ui.table_all_questions.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.ResizeToContents)
        self.ui.table_exam_questions.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.ResizeToContents)

        # (Opsiyonel) Eğer son sütunun ekrana yapışmasını sağlayan bir ayar açıksa onu kapat:
        self.ui.table_all_questions.horizontalHeader().setStretchLastSection(False)
        self.ui.table_exam_questions.horizontalHeader().setStretchLastSection(False)

    def yeni_sinav_alani_degistir(self):
        """+ butonuna basıldığında yeni sınav ekleme kutularını açar/kapatır."""
        su_anki_durum = self.ui.txt_yeni_sinav_adi.isVisible()

        # Eğer görünürse gizle, gizliyse görünür yap (Toggle mantığı)
        self.ui.txt_yeni_sinav_adi.setVisible(not su_anki_durum)
        self.ui.btn_sinav_kaydet.setVisible(not su_anki_durum)
        self.ui.lbl_yeni_sinav.setVisible(not su_anki_durum)
        # Kutu açıldığında kullanıcının imleci otomatik olarak içine girsin
        if not su_anki_durum:
            self.ui.txt_yeni_sinav_adi.setFocus()

    def yeni_sinav_kaydet(self):
        """Yeni sınavı veritabanına kaydeder ve Combobox'ı günceller."""
        sinav_adi = self.ui.txt_yeni_sinav_adi.text().strip()

        if not sinav_adi:
            QMessageBox.warning(self, "Uyarı", "Sınav adı boş olamaz!")
            return

        basarili = self.db.sinav_ekle(sinav_adi)

        # 2. Sonuca göre kullanıcıya bilgi ver ve arayüzü güncelle
        if basarili:
            QMessageBox.information(
                self, "Başarılı", f"'{sinav_adi}' başarıyla oluşturuldu.")

            self.ui.txt_yeni_sinav_adi.clear()
            self.yeni_sinav_alani_degistir()  # Kutuları tekrar gizle
            self.sinavlari_comboboxa_yukle()  # Listeyi güncelle
        else:
            QMessageBox.critical(
                self, "Hata", "Sınav oluşturulurken veritabanında sistemsel bir sorun oluştu.")

    def sinavlari_comboboxa_yukle(self):
        """Veritabanındaki sınavları Combobox'a (gizli ID'leriyle birlikte) doldurur."""
        self.ui.cb_sinavlar.clear()

        # 1. Varsayılan Seçenek (Bunun ID'sini bilerek None yapıyoruz)
        self.ui.cb_sinavlar.addItem("--- Bir Sınav Seçin ---", None)

        # 2. Veritabanından sınavları çek
        sinavlar = self.db.tum_sinavlari_getir()

        # 3. Combobox'a Ekleme Döngüsü
        for sinav in sinavlar:
            # İstersen ismin yanına tarih de koyabilirsin: f"{sinav['ad']} ({sinav['tarih']})"
            gosterilecek_metin = sinav['ad']

            # addItem(Görünen_Metin, Gizli_Veri) mantığıyla çalışır
            self.ui.cb_sinavlar.addItem(gosterilecek_metin, sinav['id'])

    def secili_sinav_sorularini_yukle(self, index):
        """Seçilen sınava ait soruları sağdaki tabloya yükler."""
        if index <= 0:
            self.ui.lbl_sinav_sorulari.setText("Sınav Soruları")
        else:
            # 1. ComboBox'tan seçili sınavın ID'sini al (UserRole'e saklamıştık)
            sinav_id = self.ui.cb_sinavlar.currentData()
            sinav_adi = self.ui.cb_sinavlar.currentText()

            # 2. Sağdaki tabloyu her seçimde temizle
            self.ui.table_exam_questions.setRowCount(0)

            # Eğer "Sınav Seçin" gibi boş bir seçenek seçiliyse (ID None ise) dur
            if sinav_id is None:
                return

            # 3. Veritabanından o sınavın sorularını çek
            sorular = self.db.sinav_sorularini_getir(sinav_id)

            # 4. Soruları sağ tabloya (Sınav Soruları) bas
            for row_idx, soru in enumerate(sorular):
                self.ui.table_exam_questions.insertRow(row_idx)

                # ID hücresi (Gizli)
                id_item = QTableWidgetItem(str(soru[0]))
                id_item.setData(Qt.ItemDataRole.UserRole, soru[0])

                # Metin hücresi
                metin_item = QTableWidgetItem(str(soru[1]))

                self.ui.table_exam_questions.setItem(row_idx, 0, id_item)
                self.ui.table_exam_questions.setItem(row_idx, 1, metin_item)

            self.ui.lbl_sinav_sorulari.setText(
                "Sınav Soruları (" + sinav_adi + ")")

            self.ui.table_exam_questions.resizeColumnsToContents()

    def secili_konu_idlerini_al(self):
        """Kendi ağacımızdaki (tree_filtre) işaretlenmiş konuların ID'lerini toplar."""
        secili_id_listesi = []

        # Ağacın en tepe noktasını (görünmez kök) alıyoruz
        root = self.ui.tree_filtre.invisibleRootItem()

        if not root:
            return []

        # 1. Seviye: Dersleri dön
        for i in range(root.childCount()):
            ders_item = root.child(i)

            # 2. Seviye: Bu dersin altındaki konuları dön
            for j in range(ders_item.childCount()):
                konu_item = ders_item.child(j)

                # Eğer konunun yanındaki kutucuk işaretliyse
                if konu_item.checkState(0) == Qt.CheckState.Checked:
                    # Gizli bölmedeki (UserRole) ID'yi al ve listeye ekle
                    konu_id = konu_item.data(0, Qt.ItemDataRole.UserRole)
                    secili_id_listesi.append(konu_id)

        return secili_id_listesi

    def tabloyu_guncelle(self, item=None, column=0):

        # Kendi ağacımızdan ID'leri topla
        secili_id_listesi = self.secili_konu_idlerini_al()

        self.ui.table_all_questions.setSortingEnabled(False)
        self.ui.table_all_questions.setRowCount(0)

        self.ui.table_all_questions.setColumnCount(2)
        self.ui.table_all_questions.setHorizontalHeaderLabels(
            ["ID", "Soru Metni"])
        self.ui.table_all_questions.setColumnHidden(0, True)

        if not secili_id_listesi:

            return

        # Veritabanına istek atıyoruz

        sorular = self.db.sorulari_getir_filtreli(secili_id_listesi)

        # Gelen verileri tabloya bas
        for row_idx, soru in enumerate(sorular):
            self.ui.table_all_questions.insertRow(row_idx)

            soru_id = soru[0]
            orijinal_metin = str(soru[1])
            resim_adi = soru[4]

            # Sütun 0: Soru ID'si
            id_item = QTableWidgetItem(str(soru_id))
            id_item.setData(Qt.ItemDataRole.UserRole, soru_id)
            id_item.setFlags(id_item.flags() & ~Qt.ItemFlag.ItemIsEditable)

            # Sütun 1: Soru Metninin Özeti
            ozet_metin = orijinal_metin[:60] + \
                "..." if len(orijinal_metin) > 60 else orijinal_metin
            if resim_adi:
                ozet_metin = "📷 " + ozet_metin

            metin_item = QTableWidgetItem(ozet_metin)
            metin_item.setFlags(metin_item.flags() & ~
                                Qt.ItemFlag.ItemIsEditable)

            self.ui.table_all_questions.setItem(row_idx, 0, id_item)
            self.ui.table_all_questions.setItem(row_idx, 1, metin_item)
        self.ui.table_all_questions.setSortingEnabled(True)
        self.ui.table_all_questions.resizeColumnsToContents()

    def havuz_satir_tiklandi(self, row, column):
        """Hangi tabloya tıklandıysa o tablodaki soruyu önizleme alanına yansıtır."""

        # 1. Sinyali gönderen (tıklanan) tabloyu dinamik olarak yakala
        tiklanan_tablo = self.sender()

        # Güvenlik kontrolü: Eğer tıklanan tablo bir QTableWidget değilse işlemi durdur
        if not tiklanan_tablo:
            return

        # 2. Artık sabit tablo adı yerine, tıklanan tabloyu (tiklanan_tablo) kullanıyoruz
        soru_id_item = tiklanan_tablo.item(row, 0)
        if not soru_id_item:
            return

        # (İpucu: Eğer ID'yi daha önce konuştuğumuz gibi UserRole içine gizlediysen,
        # burayı int(soru_id_item.data(Qt.ItemDataRole.UserRole)) olarak değiştirebilirsin.
        # Ekranda görünüyorsa text() kalabilir.)
        soru_id = int(soru_id_item.text())

        # 3. Sadece ID'yi bildiğimiz için sorunun tüm detaylarını veritabanından çekiyoruz
        soru_data, _ = self.db.soru_detay_getir(soru_id)

        if soru_data:
            # --- METİN VE ŞIKLARI FORMATLAMA ---
            soru_metni = soru_data.get('metin', '')
            dogru_cevap = soru_data.get('dogru_cevap', '')
            siklar_raw = soru_data.get('siklar_json', '[]')

            siklar_text = ""
            try:
                siklar_liste = json.loads(siklar_raw)
                siklar_text = "\n".join(siklar_liste)
            except:
                siklar_text = str(siklar_raw)

            formatli_siklar = []
            for harf, metin in siklar_liste.items():
                # Harfi büyütüp yanına parantez ekliyoruz. Örn: "A) asdsad"
                satir = f"{harf}) {metin}"
                formatli_siklar.append(satir)

            # 3. Hazırladığımız bu yeni listeyi alt alta birleştiriyoruz
            siklar_text = "\n".join(formatli_siklar)
            tam_metin = f"{soru_metni}\n\n{siklar_text}\n\nDoğru Cevap: {dogru_cevap}"

            self.ui.txt_onizleme.setText(tam_metin)

            # --- RESİM YÜKLEME İŞLEMİ ---
            self.ui.lbl_resim_sinav.clear()
            self.ui.lbl_resim_sinav.setText("Görsel Yok")

            resim_adi = soru_data.get('resim_adi')
            if resim_adi:

                appdata_klasoru = Path(
                    os.getenv('LOCALAPPDATA')) / "SoruBankasi" / "resimler"
                resim_yolu = appdata_klasoru / resim_adi

                if resim_yolu.exists():
                    self.onizlenen_resim_yolu = str(resim_yolu)
                    pixmap = QPixmap(str(resim_yolu))
                    self.ui.lbl_resim_sinav.setMinimumSize(150, 150)
                    sigdirilmis_pixmap = pixmap.scaled(
                        self.ui.lbl_resim_sinav.size(),
                        Qt.AspectRatioMode.KeepAspectRatio,
                        Qt.TransformationMode.SmoothTransformation
                    )
                    self.ui.lbl_resim_sinav.setPixmap(sigdirilmis_pixmap)
                else:
                    self.onizlenen_resim_yolu = None
                    self.ui.lbl_resim_sinav.setText("⚠️ Görsel dosyası kayıp!")

    def sinava_soru_ekle(self):
        """Seçili soruyu/soruları havuzdan alıp Sınav Soruları tablosuna kopyalar."""

        # -------------------------------------------------------------------
        # Gizli sütun (ID) selectedIndexes()'te yer almıyor.
        # Bu yüzden satır numaralarını alıp, model üzerinden ID'yi çekiyoruz.
        # -------------------------------------------------------------------
        # Sıralamayı geçici olarak kapat
        self.ui.table_all_questions.setSortingEnabled(False)

        # Hem satır hem hücre seçimini destekle: tüm seçili hücrelerden satır numaralarını topla
        secili_satirlar = set()
        for index in self.ui.table_all_questions.selectionModel().selectedIndexes():
            secili_satirlar.add(index.row())

        if not secili_satirlar:
            QMessageBox.information(
                self, "Bilgi", "Lütfen havuzdan eklenecek bir soru seçin.")
            self.ui.table_all_questions.setSortingEnabled(True)
            return

        eklenen_soru_sayisi = 0

        # Sınav tablosunda seçili satırı bul
        exam_table = self.ui.table_exam_questions
        selected_exam_rows = set()
        for idx in exam_table.selectionModel().selectedIndexes():
            selected_exam_rows.add(idx.row())
        if selected_exam_rows:
            insert_row = max(selected_exam_rows) + 1
        else:
            insert_row = exam_table.rowCount()

        for row in sorted(secili_satirlar):
            id_item = self.ui.table_all_questions.item(row, 0)
            metin_item = self.ui.table_all_questions.item(row, 1)

            if not id_item or not metin_item:
                continue
            soru_id = int(id_item.data(Qt.ItemDataRole.UserRole))
            soru_metin = metin_item.text()

            # 3. KONTROL: Bu soru zaten sağdaki tabloda (Sınavda) var mı?
            zaten_var = False
            for sag_row in range(exam_table.rowCount()):
                sag_id_item = exam_table.item(sag_row, 0)
                if sag_id_item and int(sag_id_item.data(Qt.ItemDataRole.UserRole)) == soru_id:
                    zaten_var = True
                    break
            if zaten_var:
                continue
            # 4. KOPYALAMA: Sağdaki tabloya seçili satırın altına ekle
            exam_table.insertRow(insert_row)
            yeni_id_item = QTableWidgetItem(str(soru_id))
            yeni_id_item.setData(Qt.ItemDataRole.UserRole, soru_id)
            yeni_id_item.setFlags(Qt.ItemFlag.ItemIsEnabled |
                                  Qt.ItemFlag.ItemIsSelectable)
            yeni_metin_item = QTableWidgetItem(soru_metin)
            yeni_metin_item.setFlags(
                Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable)
            exam_table.setItem(insert_row, 0, yeni_id_item)
            exam_table.setItem(insert_row, 1, yeni_metin_item)
            insert_row += 1
            eklenen_soru_sayisi += 1

        self.ui.table_all_questions.clearSelection()
        # Sıralamayı tekrar aç
        self.ui.table_all_questions.setSortingEnabled(True)

    def sinavdan_soru_cikar(self):
        """Sağdaki sınav tablosundan seçili soruları taslaktan siler."""

        secili_satirlar = set()
        for item in self.ui.table_exam_questions.selectedItems():
            secili_satirlar.add(item.row())

        if not secili_satirlar:
            return

        # Çok Önemli: Tablodan satır silerken her zaman "Aşağıdan Yukarıya (Tersten)" silinmelidir!
        # Aksi takdirde indeksler kayar ve program çöker.
        for row in sorted(secili_satirlar, reverse=True):
            self.ui.table_exam_questions.removeRow(row)

    def sinavi_tamamen_kaydet(self):
        """Sağ tablodaki tüm soruları seçili sınava kalıcı olarak yazar."""

        # 1. Hangi sınav seçili?
        sinav_id = self.ui.cb_sinavlar.currentData()
        if sinav_id is None:
            QMessageBox.warning(self, "Uyarı", "Lütfen önce bir sınav seçin!")
            return

        # 2. Sağdaki tabloda (Sınav Soruları) hangi soru ID'leri var?
        soru_idleri = []
        for row in range(self.ui.table_exam_questions.rowCount()):
            id_item = self.ui.table_exam_questions.item(row, 0)
            if id_item:
                soru_idleri.append(int(id_item.text()))

        if not soru_idleri:
            confirm = QMessageBox.question(self, "Emin misiniz?",
                                           "Sınavda hiç soru yok. Tüm sorular silinsin mi?",
                                           QMessageBox.Yes | QMessageBox.No)
            if confirm == QMessageBox.No:
                return

        # 3. Veritabanına toplu kayıt isteği gönder
        basarili = self.db.sinav_sorularini_kaydet(sinav_id, soru_idleri)

        if basarili:
            QMessageBox.information(
                self, "Başarılı", "Sınav soruları başarıyla güncellendi!")
        else:
            QMessageBox.critical(
                self, "Hata", "Kaydedilirken bir sorun oluştu.")

    def soru_kopyala(self):
        """Soru metnini ve görselini Word'e tam uyumlu HTML formatında kopyalar."""

        # 1. Metni Al ve Hazırla
        # Arayüzdeki düz metni alıyoruz
        soru_metni = self.ui.txt_onizleme.toPlainText()

        # \n (Enter) karakterlerini HTML satır atlama etiketine (<br>) çeviriyoruz
        html_satirlar = soru_metni.replace('\n', '<br>')

        # 2. Resmi Al ve Base64'e Çevir
        img_html = ""

        if self.onizlenen_resim_yolu:
            # Resmi belleğe yazıp Base64 formatına kodluyoruz
            orijinal_pixmap = QPixmap(self.onizlenen_resim_yolu)

            byte_array = QByteArray()
            buffer = QBuffer(byte_array)
            buffer.open(QIODevice.OpenModeFlag.WriteOnly)

            # Tam kaliteli resmi PNG olarak belleğe yaz
            orijinal_pixmap.save(buffer, "PNG", quality=100)
            base64_data = byte_array.toBase64().data().decode('utf-8')

            # Max-width ile Word'de sayfa dışına taşmasını engelliyoruz,
            # ancak resmin pikselleri orijinal kalitesiyle korunuyor!
            img_html = f'<br><br><img src="data:image/png;base64,{base64_data}" width="200">'

        # 3. Senin Şablonun ile Değişkenleri Birleştirme
        html_icerik = f"""
            <html>
            <head>
            <style>
                /* Tüm gövdeyi Times New Roman yap ve satır yüksekliklerini sıfırla */
                body {{
                    font-family: 'Times New Roman', Times, serif;
                    font-size: 12pt;  /* Word standardı genelde 11 veya 12pt'dir */
                    line-height: 1.0; /* Satır aralığı: Tek (Single) */
                    margin: 0;
                    padding: 0;
                }}
                /* Paragraf boşluklarını tamamen ez */
                p {{
                    margin-top: 0pt;
                    margin-bottom: 0pt;
                    padding: 0;
                }}
            </style>
            </head>
            <body>
                {img_html}
                <p>{html_satirlar}</p>
            </body>
            </html>
        """

        # 4. Panoya (Clipboard) Yükleme
        clipboard = QApplication.clipboard()
        mime_data = QMimeData()

        # Hem Word'ün anlayacağı HTML formatını hem de Not Defteri için düz metni yüklüyoruz
        mime_data.setHtml(html_icerik)
        mime_data.setText(soru_metni)

        clipboard.setMimeData(mime_data)

        # Kontrol için konsola yazdır
      

    def word_dosyasi_olustur(self):
        """ComboBox'ta seçili sınava ait tüm soruları veritabanından çeker ve Word belgesi üretir."""

        # ==========================================
        # 1. COMBOBOX KONTROLÜ VE SINAV BİLGİSİ
        # ==========================================
        # Not: ComboBox nesnenin adını kendi arayüzüne göre güncelle (örn: cmb_sinavlar)
        combo_box = self.ui.cb_sinavlar

        # Eğer "--- Bir Sınav Seçin ---" gibi 0. indeks seçiliyse işlemi durdur
        if combo_box.currentIndex() <= 0:
            QMessageBox.warning(
                self, "Uyarı", "Lütfen Word dosyası oluşturmak için bir sınav seçin!")
            return

        sinav_adi = combo_box.currentText()
        sinav_id = combo_box.currentData()

        # PROFESYONEL İPUCU: Eğer sınav ID'sini ComboBox'ın UserRole'üne gizlediysen
        # sinav_id = combo_box.currentData() şeklinde çekmek isimden aramaktan çok daha güvenlidir.

        # ==========================================
        # 2. VERİTABANINDAN SORULARI ÇEKME
        # ==========================================
        # db sınıfında yazacağın bu fonksiyonun, o sınava ait soruları bir liste içinde
        # sözlük (dictionary) olarak döndürdüğünü varsayıyoruz.
        sorular = self.db.sinav_sorularini_getir(sinav_id)

        if not sorular:
            QMessageBox.warning(
                self, "Uyarı", f"'{sinav_adi}' adlı sınavda henüz hiç soru yok!")
            return

        # ==========================================
        # 3. KAYIT YERİ SEÇİMİ
        # ==========================================
        dosya_yolu, _ = QFileDialog.getSaveFileName(
            self,
            "Word Dosyasını Kaydet",
            # Sınav adını varsayılan dosya adı yap
            os.path.expanduser(
                f"~/Desktop/{sinav_adi.replace(' ', '_')}.docx"),
            "Word Belgesi (*.docx)"
        )

        if not dosya_yolu:
            return

        # ==========================================
        # 4. WORD BELGESİNİ OLUŞTURMA
        # ==========================================
        try:
            doc = Document()

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Bazı Word sürümlerinde fontun tam uygulanması için gerekli alt ayar
            style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            style.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')

            # Sınav Başlığını Ekle
            baslik = doc.add_heading(sinav_adi.upper(), 0)
            baslik.alignment = 1  # Ortala

            for run in baslik.runs:
                run.font.name = 'Times New Roman'
                # Başlık biraz daha büyük kalabilir (isteğe bağlı)
                run.font.size = Pt(16)
                run.bold = True

            # Veritabanından gelen her bir soru için döngü oluştur
            for index, soru_data in enumerate(sorular):
                soru_numarasi = index + 1

                # --- Soru Metni ---
                soru_data = dict(soru_data)

                soru_metni = soru_data.get('metin', '')
                paragraf = doc.add_paragraph()
                paragraf.add_run(f"Soru {soru_numarasi}: ").bold = True
                paragraf.add_run(soru_metni)

                # --- Resim İşlemi ---
                resim_adi = soru_data.get('resim_adi')
                if resim_adi:
                    resim_tam_yol = Path(
                        os.getenv('LOCALAPPDATA')) / "SoruBankasi" / "resimler" / resim_adi
                    if resim_tam_yol.exists():
                        doc.add_picture(str(resim_tam_yol), width=Inches(2.5))

                # --- Şıklar ---
                siklar_raw = soru_data.get('siklar_json', '[]')
                try:
                    siklar_dict = json.loads(siklar_raw)
                    for harf, metin in siklar_dict.items():
                        sik_paragrafi = doc.add_paragraph()

                        # Şıklar arasındaki boşluğu da sıfırlıyoruz
                        sik_paragrafi.paragraph_format.space_before = Pt(0)
                        sik_paragrafi.paragraph_format.space_after = Pt(0)
                        sik_paragrafi.paragraph_format.line_spacing = 1.0

                        sik_run = sik_paragrafi.add_run(
                            f"{harf.upper()}) {metin}")
                        sik_run.font.name = 'Times New Roman'
                        sik_run.font.size = Pt(12)
                        sik_paragrafi.paragraph_format.left_indent = Inches(
                            0.2)

                except:
                    hata_paragraf = doc.add_paragraph(str(siklar_raw))
                    hata_paragraf.paragraph_format.space_after = Pt(0)
                bosluk_paragrafi = doc.add_paragraph()
                # Bu boşluk paragrafının boyutunu 12 punto yaparak standart bir boşluk bırakıyoruz
                bosluk_paragrafi.paragraph_format.space_before = Pt(12)
                bosluk_paragrafi.paragraph_format.space_after = Pt(0)
            # Dosyayı Kaydet
            doc.save(dosya_yolu)
            QMessageBox.information(
                self, "Başarılı", f"{len(sorular)} soruluk sınav başarıyla Word'e aktarıldı!")

        except PermissionError:
            QMessageBox.critical(
                self, "Hata", "Dosya başka bir programda açık. Lütfen Word'ü kapatıp tekrar deneyin.")
        except Exception as e:
            QMessageBox.critical(
                self, "Hata", f"Word dosyası oluşturulurken hata oluştu:\n{str(e)}")

    def sinav_sil(self):
        """Seçili sınavı kullanıcı onayıyla siler ve arayüzü günceller."""
        
        # 1. ComboBox'tan ID'yi al
        sinav_id = self.ui.cb_sinavlar.currentData()
        sinav_adi = self.ui.cb_sinavlar.currentText()

        # Eğer geçerli bir seçim yoksa (None ise) işlemi durdur
        if sinav_id is None:
            QMessageBox.warning(self, "Uyarı", "Lütfen silmek istediğiniz sınavı seçin.")
            return

        # 2. Türkçe Butonlu Onay Penceresi
        onay_box = QMessageBox(self)
        onay_box.setWindowTitle("Sınavı Sil")
        onay_box.setText(f"'{sinav_adi}' isimli sınavı silmek istediğinize emin misiniz?\n\nBu işlem geri alınamaz.")
        onay_box.setIcon(QMessageBox.Icon.Question)
        
        evet_btn = onay_box.addButton("Evet", QMessageBox.ButtonRole.YesRole)
        hayir_btn = onay_box.addButton("Hayır", QMessageBox.ButtonRole.NoRole)
        onay_box.setDefaultButton(hayir_btn)
        
        onay_box.exec()

        if onay_box.clickedButton() != evet_btn:
            return # Kullanıcı Hayır dedi

        # 3. Veritabanından Silme İşlemini Başlat
        basarili = self.db.sinav_veritabanindan_sil(sinav_id)

        if basarili:
            QMessageBox.information(self, "Başarılı", f"'{sinav_adi}' sınavı başarıyla silindi.")
            
            # 4. ARAYÜZÜ GÜNCELLE (Kritik Adım)
            # Sınavlar listesini (ComboBox) yeniden yükle ki silinen sınav gitsin
            self.sinavlari_comboboxa_yukle() 
            
            # Eğer sağ tarafta o sınava ait sorular listeleniyorsa tabloyu da temizle
            self.ui.table_exam_questions.setRowCount(0)
        else:
            QMessageBox.critical(self, "Hata", "Sınav silinirken teknik bir sorun oluştu.")

